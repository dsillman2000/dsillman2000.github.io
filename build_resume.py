#!/usr/bin/env python3
"""Generate David Sillman's resume as a .docx with Swimmer Light styling."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BLUE = RGBColor(0x2B, 0x6C, 0xB0)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE_HEX = "2B6CB0"
DARK_HEX = "1A1A1A"
GRAY_HEX = "666666"
BORDER_GRAY_HEX = "CCCCCC"

FONT = "Inter"
MONO = "Roboto Mono"

OUTPUT_DIR = os.environ.get("OUTPUT_DIR", os.getcwd())
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "David_Sillman_Resume.docx")


def styled_run(paragraph, text, font_name=FONT, size=Pt(10), color=DARK,
               bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.insert(0, rFonts)
    return run


def para_spacing(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after = Twips(after)
    if line is not None:
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
        pPr.append(sp)


def add_bottom_border(p, color=BLUE_HEX, sz="6"):
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url, font_name=MONO, size=Pt(8),
                  color_hex=BLUE_HEX):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(szCs)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def section_heading(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=240, after=80)
    run = styled_run(p, text.upper(), size=Pt(10), color=BLUE, bold=True)
    rPr = run._element.get_or_add_rPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:val"), "30")
    rPr.append(spacing_el)
    add_bottom_border(p, color=BLUE_HEX, sz="6")
    return p


def parse_bold_segments(text):
    """Return list of (text, is_bold) from markdown **bold** markup."""
    segments = []
    while "**" in text:
        idx = text.index("**")
        if idx > 0:
            segments.append((text[:idx], False))
        text = text[idx + 2:]
        end = text.index("**")
        segments.append((text[:end], True))
        text = text[end + 2:]
    if text:
        segments.append((text, False))
    return segments


def add_bullet(doc, text, indent_emu=Inches(0.35), hanging_emu=Inches(0.2),
               size=Pt(10)):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=40)
    pf = p.paragraph_format
    pf.left_indent = indent_emu
    pf.first_line_indent = -hanging_emu

    styled_run(p, "\u2022  ", size=size, color=DARK)

    segments = parse_bold_segments(text)
    for seg_text, is_bold in segments:
        styled_run(p, seg_text, size=size, color=DARK, bold=is_bold)
    return p


def add_sub_bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=40)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.7)
    pf.first_line_indent = Inches(-0.2)

    styled_run(p, "\u2013  ", size=size, color=GRAY)

    segments = parse_bold_segments(text)
    for seg_text, is_bold in segments:
        styled_run(p, seg_text, size=size, color=DARK, bold=is_bold)
    return p


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for attr_name, attr_val in val.items():
            el.set(qn(f"w:{attr_name}"), attr_val)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def build_resume():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(8)
    style.font.color.rgb = DARK
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rPr.insert(0, rFonts)

    # =========================================================================
    # HEADER
    # =========================================================================
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_name, before=0, after=25)
    styled_run(p_name, "David Sillman", size=Pt(20), color=BLUE, bold=True)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_contact, before=0, after=10)
    styled_run(p_contact, "Woburn, MA  |  (650) 933-8448  |  ", size=Pt(8),
               color=GRAY)
    styled_run(p_contact, "dsillman2000@gmail.com", font_name=MONO,
               size=Pt(7.5), color=BLUE)

    p_links = doc.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_links, before=0, after=30)

    add_hyperlink(p_links, "LinkedIn",
                  "https://www.linkedin.com/in/david-sillman-8abb5831")
    styled_run(p_links, "   |   ", size=Pt(8), color=GRAY)
    add_hyperlink(p_links, "GitHub", "https://github.com/dsillman2000")
    styled_run(p_links, "   |   ", size=Pt(8), color=GRAY)
    add_hyperlink(p_links, "Personal Website", "https://www.dsillman.com")

    p_rule = doc.add_paragraph()
    para_spacing(p_rule, before=0, after=0)
    add_bottom_border(p_rule, color=BORDER_GRAY_HEX, sz="4")

    # =========================================================================
    # PROFESSIONAL SUMMARY
    # =========================================================================
    section_heading(doc, "Professional Summary")

    summary = (
        "Senior Data Engineer with 4+ years architecting data pipelines and "
        "cloud infrastructure at scale. Expert in dbt, Snowflake, and Airflow "
        "with deep Terraform/OpenTofu proficiency on AWS and Cloudflare. "
        "Bridges analytics and software engineering, relishing the challenges "
        "of designing cloud-native solutions for clinical and commercial "
        "reporting while championing CI/CD, IaC, and agentic AI tooling. "
        "Strong foundation in Python, Rust, and SQL."
    )
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=30)
    styled_run(p, summary, size=Pt(8), color=DARK)

    # =========================================================================
    # TECHNICAL SKILLS
    # =========================================================================
    section_heading(doc, "Technical Skills")

    skills_data = [
        ("Languages", "Python, Rust, SQL, TypeScript, R"),
        ("Data & Pipelines", "dbt, Snowflake, Databricks, Apache Airflow, PySpark, Polars"),
        ("Cloud & IaC", "AWS, Cloudflare, Terraform / OpenTofu, Ansible"),
        ("Platform & DevOps", "HashiCorp Vault, JFrog Artifactory, GitLab CI/CD, GitHub Actions, Docker"),
        ("Web & Visualization", "FastAPI, Plotly, Streamlit, Tableau, React"),
        ("AI & Automation", "Agentic AI tooling, custom ACP integrations"),
    ]

    table = doc.add_table(rows=len(skills_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    remove_table_borders(table)

    col_widths = [Inches(1.5), Inches(5.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    hdr = table.rows[0]
    for i, text in enumerate(["Category", "Technologies"]):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        para_spacing(p, before=10, after=10)
        styled_run(p, text, size=Pt(8), color=BLUE, bold=True)
        set_cell_border(cell, bottom={
            "val": "single", "sz": "6", "color": BLUE_HEX, "space": "0"
        })

    for row_idx, (cat, techs) in enumerate(skills_data):
        row = table.rows[row_idx + 1]
        cell_cat = row.cells[0]
        cell_cat.paragraphs[0].clear()
        p = cell_cat.paragraphs[0]
        para_spacing(p, before=10, after=10)
        styled_run(p, cat, size=Pt(8), color=DARK, bold=True)
        cell_cat.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell_tech = row.cells[1]
        cell_tech.paragraphs[0].clear()
        p = cell_tech.paragraphs[0]
        para_spacing(p, before=10, after=10)
        styled_run(p, techs, size=Pt(8), color=DARK)
        cell_tech.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for cell in row.cells:
            set_cell_border(cell, bottom={
                "val": "single", "sz": "2", "color": "E5E7EB", "space": "0"
            })

    # =========================================================================
    # PROFESSIONAL EXPERIENCE
    # =========================================================================
    section_heading(doc, "Professional Experience")

    p = doc.add_paragraph()
    para_spacing(p, before=80, after=10)
    styled_run(p, "Intuitive Surgical", size=Pt(9.5), color=DARK, bold=True)
    styled_run(p, "   |   ", size=Pt(8.5), color=GRAY)
    styled_run(p, "Sunnyvale, CA", size=Pt(8.5), color=GRAY, italic=True)

    p = doc.add_paragraph()
    para_spacing(p, before=25, after=30)
    styled_run(p, "Senior Data Engineer", size=Pt(8.5), color=BLUE, bold=True)
    styled_run(p, "   |   ", size=Pt(8), color=GRAY)
    styled_run(p, "July 2022 \u2013 Present", size=Pt(8), color=GRAY,
               italic=True)

    sde_bullets = [
        "**Architected Data Pipelines:** Deployed cloud-based pipelines for "
        "clinical and commercial analytics using **dbt**, **Snowflake**, and "
        "**Databricks**.",

        "**Orchestrated Workflows:** Built periodic data collection tools "
        "with **Python** and **Apache Airflow** for reliable downstream "
        "ingestion.",

        "**Managed Cloud Infrastructure:** Owned **Terraform/OpenTofu** on "
        "**AWS**, administering **HashiCorp Vault** and **JFrog Artifactory** "
        "for a multi-tenant organization; uses **Ansible** for onboarding "
        "automation.",

        "**Led Agentic AI Integration:** Drove adoption of agentic AI tools "
        "on the Ion data team, authoring shared skills and scripts via the "
        "company\u2019s custom **ACP**.",

        "**High-Speed Data Processor:** Built a multi-threaded, **GB/s** file "
        "processor in **Rust**, eliminating ingestion bottlenecks.",

        "**Engineering Best Practices:** Enforced **automated CI/CD** testing "
        "and documentation in **GitLab** (work) and **GitHub Actions** "
        "(personal), improving code quality and velocity.",

        "**Optimized Data Processing:** Migrated in-memory Python scripts to "
        "**SQL** in cloud warehouses and **PySpark** for lazy out-of-memory "
        "operations.",
    ]
    for bullet in sde_bullets:
        add_bullet(doc, bullet, size=Pt(8))

    p = doc.add_paragraph()
    para_spacing(p, before=65, after=30)
    styled_run(p, "Advanced Data Analyst", size=Pt(8.5), color=BLUE,
               bold=True)
    styled_run(p, "   |   ", size=Pt(8), color=GRAY)
    styled_run(p, "July 2022 \u2013 March 2023  ", size=Pt(8), color=GRAY,
               italic=True)
    styled_run(p, "(concurrent)", size=Pt(7.5), color=GRAY, italic=True)

    ada_bullets = [
        "Analyzed clinical usage data for the Ion System with **Python**, "
        "**Pandas**, and **MSSQL**.",

        "Built internal web apps: a **Dash/Plotly** visualization dashboard "
        "with **Jinja** and a **React.js/VTK.js** 3D geometry viewer.",
    ]
    for bullet in ada_bullets:
        add_bullet(doc, bullet, size=Pt(8))

    p = doc.add_paragraph()
    para_spacing(p, before=65, after=5)
    styled_run(p, "Data Science Intern", size=Pt(8.5), color=BLUE, bold=True)
    styled_run(p, "   |   ", size=Pt(8), color=GRAY)
    styled_run(p, "May 2020 \u2013 April 2022", size=Pt(8), color=GRAY,
               italic=True)

    add_bullet(
        doc,
        "Analyzed clinical usage data, laying the groundwork for the "
        "pipelines and apps developed in subsequent roles.",
        size=Pt(8),
    )

    # =========================================================================
    # SELECT PROJECT HIGHLIGHTS
    # =========================================================================
    section_heading(doc, "Select Project Highlights")

    p = doc.add_paragraph()
    para_spacing(p, before=25, after=10)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.35)
    pf.first_line_indent = Inches(-0.2)
    styled_run(p, "\u2022  ", size=Pt(8), color=DARK)
    styled_run(p, "decision-tree-stuff", font_name=MONO, size=Pt(7.5),
               color=BLUE, bold=True)
    styled_run(
        p,
        ": Binary decision tree classifier supporting lazy and eager "
        "Polars DataFrames in Python.",
        size=Pt(8), color=DARK,
    )

    p = doc.add_paragraph()
    para_spacing(p, before=0, after=10)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.35)
    pf.first_line_indent = Inches(-0.2)
    styled_run(p, "\u2022  ", size=Pt(8), color=DARK)
    styled_run(p, "yaml-reference-*", font_name=MONO, size=Pt(7.5),
               color=BLUE, bold=True)
    styled_run(
        p,
        ": Python/TypeScript library for resolving YAML documents with "
        "custom referential tags; test-driven, multi-language.",
        size=Pt(8), color=DARK,
    )

    add_bullet(
        doc,
        "Various **Terraform** and **Cloudflare** projects on GitHub, "
        "including IaC-managed Workers, DNS, and tunnel deployments for "
        "personal cloud infrastructure.",
        size=Pt(8),
    )

    # =========================================================================
    # EDUCATION
    # =========================================================================
    section_heading(doc, "Education")

    p = doc.add_paragraph()
    para_spacing(p, before=50, after=10)
    styled_run(p, "Purdue University", size=Pt(9.5), color=DARK, bold=True)
    styled_run(p, "   |   ", size=Pt(8.5), color=GRAY)
    styled_run(p, "West Lafayette, IN", size=Pt(8.5), color=GRAY, italic=True)

    p = doc.add_paragraph()
    para_spacing(p, before=0, after=0)
    styled_run(
        p,
        "Bachelor of Science (B.S.) in Data Science and Mathematics "
        "(Double Major)",
        size=Pt(8.5), color=DARK, bold=True,
    )
    styled_run(p, "   |   ", size=Pt(8.5), color=GRAY)
    styled_run(p, "August 2018 \u2013 May 2022", size=Pt(8.5), color=GRAY,
               italic=True)

    # =========================================================================
    # SAVE
    # =========================================================================
    doc.save(OUTPUT_FILE)
    print(f"Resume saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_resume()
